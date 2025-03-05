# import os 
# import fitz  
# import markdown 
# import chromadb 
# from sentence_transformers import SentenceTransformer 
# import docx  
# import numpy as np
# import logging

# # Configure logging
# logging.basicConfig(level=logging.INFO, 
#                     format='%(asctime)s - %(levelname)s: %(message)s')

# # Initialize embedding model 
# embedding_model = SentenceTransformer("all-MiniLM-L6-v2") 
 
# # Initialize ChromaDB 
# chroma_client = chromadb.PersistentClient(path="./data_store") 
# collection = chroma_client.get_or_create_collection(
#     name="documents", 
#     metadata={"hnsw:space": "cosine"}
# ) 

# def safe_float_conversion(embedding):
#     """
#     Safely convert embedding to list of floats.
#     Handles various input types robustly.
#     """
#     try:
#         # If it's a numpy array, convert to list
#         if hasattr(embedding, 'tolist'):
#             embedding = embedding.tolist()
        
#         # Ensure each element is a float
#         return [float(x) for x in embedding]
#     except Exception as e:
#         logging.error(f"Embedding conversion error: {e}")
#         return None

# def store_in_vector_db(file_path):
#     """Processes a file, generates embeddings, and stores in ChromaDB with robust error handling."""
#     file_name = os.path.basename(file_path)

#     try:
#         # Extract text based on file type
#         if file_path.endswith(".pdf"): 
#             text = extract_text_from_pdf(file_path) 
#         elif file_path.endswith(".md"): 
#             text = extract_text_from_markdown(file_path) 
#         elif file_path.endswith('.docx') or file_path.endswith('.doc'): 
#             text = extract_text_from_word(file_path) 
#         else: 
#             logging.warning(f"Skipping unsupported file: {file_name}") 
#             return 

#         # Chunk text 
#         chunks = chunk_text(text) 

#         logging.info(f"Processing {file_name}: {len(chunks)} chunks generated")

#         # Generate embeddings with error handling
#         try:
#             embeddings = embedding_model.encode(chunks)
            
#             # Ensure embeddings are in a processable format
#             if not isinstance(embeddings, list):
#                 embeddings = embeddings.tolist()

#         except Exception as e:
#             logging.error(f"Embedding generation error for {file_name}: {e}")
#             return

#         # Store in vector DB with explicit error handling
#         successful_chunks = 0
#         for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)): 
#             try:
#                 # Safe embedding conversion
#                 safe_embedding = safe_float_conversion(embedding)
                
#                 if safe_embedding is None:
#                     logging.warning(f"Skipping chunk {idx} due to embedding conversion failure")
#                     continue

#                 # Use upsert to handle potential duplicates
#                 collection.upsert( 
#                     ids=[f"{file_name}_{idx}"], 
#                     embeddings=[safe_embedding], 
#                     metadatas=[{
#                         "source": file_name, 
#                         "chunk_id": idx, 
#                         "text": chunk
#                     }] 
#                 )
#                 successful_chunks += 1

#             except Exception as e:
#                 logging.error(f"Error storing chunk {idx} of {file_name}: {e}")

#         logging.info(f"Successfully stored {successful_chunks}/{len(chunks)} chunks from {file_name}")

#     except Exception as e:
#         logging.error(f"Critical error processing {file_name}: {e}")

# # Keep other existing functions from the original script
 
# def extract_text_from_pdf(pdf_path): 
#     """Extracts text from a PDF file.""" 
#     doc = fitz.open(pdf_path) 
#     text = "" 
#     for page in doc: 
#         text += page.get_text("text") + "\n" 
#     return text 
 
 
# def extract_text_from_markdown(md_path): 
#     """Extracts text from a Markdown file.""" 
#     with open(md_path, "r", encoding="utf-8") as f: 
#         text = f.read() 
#     return markdown.markdown(text) 
 
 
# def extract_text_from_word(word_path):
#     """Extracts text from a Word document (.docx or .doc)."""
#     try:
#         # For .docx files
#         if word_path.endswith('.docx'):
#             doc = docx.Document(word_path)
#             # Extract text from paragraphs
#             text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
#             return text
        
#         # For .doc files, you might need to use a different library like python-docx or win32com
#         # This is a placeholder - you may need to implement specific .doc handling
#         else:
#             print(f"Unsupported Word document format: {word_path}")
#             return ""
#     except Exception as e:
#         print(f"Error extracting text from {word_path}: {e}")
#         return ""
 
 
# def chunk_text(text, chunk_size=500): 
#     """Splits text into chunks of given size.""" 
#     words = text.split() 
#     return [" ".join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size)] 
 
 
# # def store_in_vector_db(file_path): 
# #     """Processes a file, generates embeddings, and stores in ChromaDB.""" 
# #     file_name = os.path.basename(file_path) 
 
# #     # Extract text 
# #     if file_path.endswith(".pdf"): 
# #         text = extract_text_from_pdf(file_path) 
# #     elif file_path.endswith(".md"): 
# #         text = extract_text_from_markdown(file_path) 
# #     elif file_path.endswith('.docx') or file_path.endswith('.doc'): 
# #         text = extract_text_from_word(file_path) 
# #     else: 
# #         print(f"Skipping unsupported file: {file_name}") 
# #         return 
 
# #     # Chunk text 
# #     chunks = chunk_text(text) 
 
# #     # Generate embeddings 
# #     embeddings = embedding_model.encode(chunks).tolist() 
 
# #     # Store in vector DB 
# #     for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)): 
# #         collection.add( 
# #             ids=[f"{file_name}_{idx}"], 
# #             embeddings=[embedding], 
# #             metadatas=[{"source": file_name, "chunk_id": idx, "text": chunk}] 
# #         ) 
 
# #     print(f"Stored {len(chunks)} chunks from {file_name} in vector database.") 
 
# # def store_in_vector_db(file_path):
# #     """Processes a file, generates embeddings, and stores in ChromaDB with detailed logging."""
# #     file_name = os.path.basename(file_path)

# #     try:
# #         # Extract text 
# #         if file_path.endswith(".pdf"): 
# #             text = extract_text_from_pdf(file_path) 
# #         elif file_path.endswith(".md"): 
# #             text = extract_text_from_markdown(file_path) 
# #         elif file_path.endswith('.docx') or file_path.endswith('.doc'): 
# #             text = extract_text_from_word(file_path) 
# #         else: 
# #             logging.warning(f"Skipping unsupported file: {file_name}") 
# #             return 

# #         # Chunk text 
# #         chunks = chunk_text(text) 

# #         logging.info(f"Processing {file_name}: {len(chunks)} chunks generated")

# #         # Generate embeddings
# #         try:
# #             embeddings = embedding_model.encode(chunks)
            
# #             # Ensure embeddings are a list of lists
# #             if not isinstance(embeddings, list):
# #                 embeddings = embeddings.tolist()

# #             logging.info(f"Embeddings generated for {file_name}")
# #             logging.info(f"Embedding shape: {len(embeddings)} x {len(embeddings[0]) if embeddings else 0}")

# #         except Exception as e:
# #             logging.error(f"Embedding generation error for {file_name}: {e}")
# #             return

# #         # Store in vector DB with explicit handling
# #         for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)): 
# #             try:
# #                 # Convert embedding to list of floats to ensure compatibility
# #                 embedding_list = [float(x) for x in embedding]
                
# #                 collection.upsert( 
# #                     ids=[f"{file_name}_{idx}"], 
# #                     embeddings=[embedding_list], 
# #                     metadatas=[{"source": file_name, "chunk_id": idx, "text": chunk}] 
# #                 )
# #             except Exception as e:
# #                 logging.error(f"Error storing chunk {idx} of {file_name}: {e}")

# #         logging.info(f"Successfully stored {len(chunks)} chunks from {file_name}")

# #     except Exception as e:
# #         logging.error(f"Error processing {file_name}: {e}")


# def process_documents(folder_path="documents"): 
#     """Processes all PDFs, Markdown, and Word files in a folder.""" 
#     for file_name in os.listdir(folder_path): 
#         file_path = os.path.join(folder_path, file_name) 
#         if file_name.endswith((".pdf", ".md", ".docx", ".doc")): 
#             store_in_vector_db(file_path) 
 
 
# if __name__ == "__main__": 
#     process_documents()

import os 
import fitz  
import markdown 
import chromadb 
from sentence_transformers import SentenceTransformer 
import docx  
import numpy as np
import logging


# Set environment variable to avoid re-downloading
os.environ["SENTENCE_TRANSFORMERS_HOME"] = "./models"

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s: %(message)s')

print("Loading model...")
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
print("Model loaded successfully!")

# Initialize ChromaDB 
chroma_client = chromadb.PersistentClient(path="./data_store") 
collection = chroma_client.get_or_create_collection(
    name="documents", 
    metadata={"hnsw:space": "cosine"}
) 

print(f"Collection name: {collection.name}")

def extract_text_from_pdf(pdf_path): 
    """Extracts text from a PDF file.""" 
    doc = fitz.open(pdf_path) 
    text = "" 
    for page in doc: 
        text += page.get_text("text") + "\n" 
    return text 

def extract_text_from_markdown(md_path): 
    """Extracts text from a Markdown file.""" 
    with open(md_path, "r", encoding="utf-8") as f: 
        text = f.read() 
    return text 

def extract_text_from_word(word_path):
    """Extracts text from a Word document."""
    try:
        if word_path.endswith('.docx'):
            doc = docx.Document(word_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            return text
        else:
            logging.warning(f"Unsupported Word document format: {word_path}")
            return ""
    except Exception as e:
        logging.error(f"Error extracting text from {word_path}: {e}")
        return ""

def chunk_text(text, chunk_size=500): 
    """Splits text into chunks of given size.""" 
    words = text.split() 
    return [" ".join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size)] 

def safe_embedding_conversion(embedding):
    """
    Safely convert embedding to a list of floats.
    Handles various input types robustly.
    """
    try:
        # Convert to list if it's a numpy array
        if hasattr(embedding, 'tolist'):
            embedding = embedding.tolist()
        
        # Ensure each element is a float
        return [float(x) for x in embedding]
    except Exception as e:
        logging.error(f"Embedding conversion error: {e}")
        return None

def store_in_vector_db(file_path):
    """Processes a file, generates embeddings, and stores in ChromaDB with comprehensive error handling."""
    file_name = os.path.basename(file_path)

    try:
        # Extract text based on file type
        if file_path.endswith(".pdf"): 
            text = extract_text_from_pdf(file_path) 
        elif file_path.endswith(".md"): 
            text = extract_text_from_markdown(file_path) 
        elif file_path.endswith('.docx'): 
            text = extract_text_from_word(file_path) 
        else: 
            logging.warning(f"Skipping unsupported file: {file_name}") 
            return 

        # Chunk text 
        chunks = chunk_text(text) 

        logging.info(f"Processing {file_name}: {len(chunks)} chunks generated")

        # Generate embeddings with comprehensive error handling
        try:
            embeddings = embedding_model.encode(chunks)
            print(f"Generated {len(embeddings)} embeddings")  # Debug

            for i, emb in enumerate(embeddings[:5]):  # Print first 5 embeddings
                print(f"Embedding {i}: {emb[:5]}...")  # Print first 5 values of each embedding

            
            # Ensure embeddings are in a processable format
            if not isinstance(embeddings, list):
                embeddings = embeddings.tolist()

        except Exception as e:
            logging.error(f"Embedding generation error for {file_name}: {e}")
            return

        # Store in vector DB with robust error handling
        successful_chunks = 0
        for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)): 
            try:
                # Safe embedding conversion
                safe_embedding = safe_embedding_conversion(embedding)
                
                if safe_embedding is None:
                    logging.warning(f"Skipping chunk {idx} due to embedding conversion failure")
                    continue

                # Ensure embedding is the correct length (384 for all-MiniLM-L6-v2)
                if len(safe_embedding) != 384:
                    logging.warning(f"Unexpected embedding length for chunk {idx}: {len(safe_embedding)}")
                    continue

                # Store the embedding
                collection.upsert( 
                    ids=[f"{file_name}_{idx}"], 
                    embeddings=[safe_embedding], 
                    metadatas=[{
                        "source": file_name, 
                        "chunk_id": idx, 
                        "text": chunk
                    }] 
                )
                successful_chunks += 1

            except Exception as e:
                logging.error(f"Error storing chunk {idx} of {file_name}: {e}")

        logging.info(f"Successfully stored {successful_chunks}/{len(chunks)} chunks from {file_name}")

    except Exception as e:
        logging.error(f"Critical error processing {file_name}: {e}")

def process_documents(folder_path="documents"):
    """Processes all supported documents in a folder."""
    files = os.listdir(folder_path)
    print(f"Files found: {files}")  # Debugging output

    for file_name in files:
        file_path = os.path.join(folder_path, file_name)
        if file_name.endswith((".pdf", ".md", ".docx")):
            print(f"Processing file: {file_name}")
            store_in_vector_db(file_path)


if __name__ == "__main__":
    # Clear existing data store if needed
    import shutil
    
    # Uncomment the following lines if you want to start with a fresh database
    # try:
    #     shutil.rmtree('./data_store', ignore_errors=True)
    # except Exception as e:
    #     print(f"Error clearing data store: {e}")
    
    process_documents()